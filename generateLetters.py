import pandas as pd
import os
from docx import Document
from datetime import datetime
import glob
import matplotlib.pyplot as plt
from docx.shared import Inches, Pt

# Flag to control testing mode
TESTING_MODE = False  # Set to False when you want real output, True for test mode

import glob
import os

def create_attendance_graph(student_name, student_attendance, output_dir, school_name, student_id):
    """
    Creates a horizontal bar graph comparing the student's attendance to the benchmark of 90% and saves it as a PNG file.
    
    Args:
    - student_name (str): The student's name.
    - student_attendance (float): The student's attendance percentage.
    - output_dir (str): The directory path where the graph image will be saved.
    - school_name (str): The name of the school.
    - student_id (int): The student's ID (used for naming the graph file).
    
    Returns:
    - str: The path to the saved graph image.
    """
    # Data for the bars
    categories = ['Target: 90%', student_name]
    attendance_values = [90, student_attendance]

    # Colors for the bars
    colors = ['green', 'red' if student_attendance < 90 else 'blue']

    # Create the horizontal bar graph
    plt.figure(figsize=(5, 1.5))  # Adjust size to be more compact (smaller height)
    
    # Plot horizontal bars
    plt.barh(categories, attendance_values, color=colors)
    
    # Set x-axis range from 0 to 100
    plt.xlim(0, 100)
    
    # Remove the title and y-axis label to save space
    plt.title('')
    plt.xlabel('')

     # Add the attendance percentage as text on the student's bar
    plt.text(attendance_values[1] - 5, 1, f"{student_attendance}%", 
             va='center', ha='center', color='white', fontweight='bold')  # Adjust position
        
    # Adjust the layout to fit the bars better
    plt.tight_layout()

    # Save the plot as a temporary image file
    graph_path = os.path.join(output_dir, f"{school_name}_{student_id}_attendance_graph.png")
    plt.savefig(graph_path)
    plt.close()

    return graph_path

def remove_old_letter_files(output_dir, school_name):
    """
    Removes old letter files in the specified output directory for a given school.
    
    Args:
    - output_dir (str): The directory path where the letter files are stored.
    - school_name (str): The name of the school (part of the file name).
    
    Returns:
    - None
    """
    # Create the pattern to match old letter files
    pattern = os.path.join(output_dir, f"{school_name}_letters_part*.docx")
    
    # Use glob to find files matching the pattern
    old_files = glob.glob(pattern)
    
    # Delete each of the old files
    for old_file in old_files:
        try:
            os.remove(old_file)
            print(f"Deleted old letter file: {old_file}")
        except OSError as e:
            print(f"Error deleting file {old_file}: {e}")

def generate_letters_by_building(results, student_languages, output_dir, building_group_mapping, csv_creation_date, max_pages=None):
    """
    Iterates through each building group, filters the results for students in that group and their respective schools,
    and passes them to the generate_student_letters function to create documents for each school within the group folder.

    Args:
    - results (dict): Dictionary of student attendance data.
    - student_languages (dict): Dictionary mapping student IDs to their preferred language.
    - output_dir (str): Directory where Word documents will be saved.
    - building_group_mapping (dict): Dictionary mapping building group names to lists of school names.
    - csv_creation_date (str): Date of the CSV file.
    - max_pages (int, optional): Maximum number of pages per document.

    Returns:
    - None
    """

    if TESTING_MODE and 'test_output' not in output_dir:
        output_dir = os.path.join(output_dir, "test_output")
        print("TESTING MODE: Saving to test_output directory.")

    for group_name, schools in building_group_mapping.items():
        for school in schools:
            # Filter students for the current school
            students_in_school = {
                student_id: student
                for student_id, student in results.items()
                if student['school_of_residence'] == school
            }

            if students_in_school:
                group_dir = os.path.join(output_dir, group_name)
                os.makedirs(group_dir, exist_ok=True)

                # Ensure we pass the `student_languages` dictionary as well
                generate_student_letters(students_in_school, student_languages, group_dir, school, csv_creation_date, max_pages)
            else:
                print(f"No students found for school {school} in group {group_name}")

def read_student_languages(language_file_path):
    """
    Reads the student language codes from a CSV file (PS_data.csv) and returns a dictionary.

    Args:
    - language_file_path (str): Path to the CSV file containing student languages.

    Returns:
    - dict: Dictionary mapping student IDs to language codes.
    """
    student_languages = {}

    try:
        df = pd.read_csv(language_file_path)
        for _, row in df.iterrows():
            student_id = str(row['student_number'])
            student_languages[student_id] = row['Written Language']
    except FileNotFoundError:
        print(f"Error: {language_file_path} not found.")
    
    return student_languages



def generate_student_letters(students_in_group, student_languages, output_dir, school_name, csv_creation_date, max_pages=None):
    """
    Generates Word documents with student letters for a single school within a group.
    Each page contains a letter for a different student, determined by their attendance and language preference.

    Args:
    - students_in_group (dict): Dictionary of students to process.
    - student_languages (dict): Dictionary mapping student IDs to their preferred language codes.
    - output_dir (str): Directory to save the documents.
    - school_name (str): Name of the school.
    - csv_creation_date (str): Date of the CSV file.
    - max_pages (int, optional): Maximum number of pages per document. Default is None.

    Returns:
    - None
    """

    # Remove old letter files for this school before generating new ones
    remove_old_letter_files(output_dir, school_name)
    
    # Initialize the document
    doc = Document()
    page_count = 0
    doc_index = 1

    for student_id, student in students_in_group.items():
        if student['current_week_att_percent'] < 90:
            # Get the student's preferred language, default to English if missing
            language_code = student_languages.get(student_id, "ENG")  # Default to English
            
            # Route to the appropriate function
            #if language_code == "SPA":
            #    generate_spanish_letter(student, doc, output_dir, school_name, csv_creation_date)
            #elif language_code == "CHK":
            #    generate_chuukese_letter(student, doc, output_dir, school_name, csv_creation_date)
            #else:
            if(language_code == "SPA"):
                generate_spanish_letter(student, doc, output_dir, school_name, csv_creation_date)
            elif(language_code == "CHK"):
                generate_chuukese_letter(student, doc, output_dir, school_name, csv_creation_date)
            else:
                generate_english_letter(student, doc, output_dir, school_name, csv_creation_date)

            # Insert a page break to separate each student's letter
            page_count += 1

            # Save and start a new document if we hit max pages
            if max_pages and page_count >= max_pages:
                doc_path = os.path.join(output_dir, f"{school_name}_letters_part{doc_index}.docx")
                doc.save(doc_path)
                print(f"Saved {doc_path}")
                
                # Reset for the next batch
                doc = Document()
                page_count = 0
                doc_index += 1

    # Save the final document if there are remaining pages
    if page_count > 0:
        doc_path = os.path.join(output_dir, f"{school_name}_letters_part{doc_index}.docx")
        doc.save(doc_path)
        print(f"Saved {doc_path}")

from docx.shared import Inches, Pt

def generate_chuukese_letter(student, doc, output_dir, school_name, csv_creation_date):
    """
    Generates a Chuukese attendance warning letter for a student with formatting adjustments 
    to fit within a single page.
    
    Args:
    - student (dict): Student information.
    - doc (Document): Word document object.
    - output_dir (str): Directory where graph images are stored.
    - school_name (str): Name of the student's school.
    - csv_creation_date (str): Date of the CSV file creation.
    
    Returns:
    - None (adds content to the provided doc object)
    """
    first_name = student['name'].split(', ')[1]
    last_name = student['name'].split(', ')[0]
    student_id = int(float(student['student_number']))  # Ensure integer format

    # Header with attendance warning in Chuukese (set font size to 13)
    header = doc.add_paragraph(f"{first_name} {last_name}, ar fifiti sukun ika attendance mi kukun seni 90% nge kemi tongeni anisi.", style='Heading 1')
    for run in header.runs:
        run.font.size = Pt(12)

    # Student Details (set font size to 10.5)
    details = doc.add_paragraph(
        f"Iten chon sukun: {student['name']}\n"
        f"Neun chon sukun ID: {student_id}\n"
        f"Sukun: {student['school_of_residence']}\n"
        f"Mwichan: {student['grade']}\n"
        f"Home Rumw: {student['home_room']}"
    )
    for run in details.runs:
        run.font.size = Pt(10.5)

    # Greeting and introduction in Chuukese (set font size to 10.5)
    greeting = doc.add_paragraph(f"Seman me inan/chon tumwunu {first_name}")
    for run in greeting.runs:
        run.font.size = Pt(10.5)
    
    intro = doc.add_paragraph(
        "Non ei putain sukun non CITY, sia achocho ngeni ach sipwe amwurinoi an ewe chon sukun attentens. "
        "Ew minichin non ranin sukun ar eto non sukun mi euchea ngeni nonomwur, ar sukun, nonomw, me memeefir. "
        "Mei kan pwaa eoch nupwen chon sukun ra etiwano ar sukun."
    )
    for run in intro.runs:
        run.font.size = Pt(10.5)

    # Attendance percentage statement in Chuukese
    attendance_statement = doc.add_paragraph(
        f"{first_name} a wor an {student['current_week_att_percent']}%, nge 90% ina met ei putai me Missouri State ra mochen epwe wor an."
    )
    for run in attendance_statement.runs:
        run.font.size = Pt(10.5)

    # Attendance Graph
    graph_path = create_attendance_graph(f"{first_name} {last_name}", student['current_week_att_percent'], output_dir, school_name, student_id)
    doc.add_picture(graph_path, width=Inches(6))
    
    try:
        os.remove(graph_path)
        print(f"Deleted temporary graph: {graph_path}")
    except OSError as e:
        print(f"Error deleting file {graph_path}: {e}")

    # Closing paragraph in Chuukese
    closing = doc.add_paragraph(
        f"Mi tongeni wate omw aninis won an {first_name} sou etiwano sukun non ei semester -kich mi fokun kinisou ren omw aninis. "
        f"Kose mochen tumwunueochu an {first_name} esap mamangeiti kunokun an sukun."
    )
    for run in closing.runs:
        run.font.size = Pt(10.5)

    # Signature
    signature = doc.add_paragraph(
        f"\nMeren,\nEwe kumien {school_name}\nSchool District"
    )
    for run in signature.runs:
        run.font.size = Pt(10.5)

    # Footer with initiative details in Chuukese (set font size to 10.5)
    footer = doc.add_paragraph(
        f"* Iei met e mak won an noum ewe chon sukun rekort ren keukun an wakareiti sukun (attendance) seni non ewe {csv_creation_date}. "
        "Ekei nampa mi pachonong an apsen mi mwumwuta me apsen ese mwumwuta. Kose mochen pusin kokori an noum ewe nenien sukun ika mi wor omw kapaseis faniten an kei apsen (etiwano ran). "
        "Ei taropwe ew an School District kinikinin angang won pekin Attendance non ei 2025, iwe sia ane ach sipwe awatetai ach weiweiti usun euchean attendance non ei School District. "
        "Kich mi esina ekoch ranin apsen (etiwano sukun) epwe faniten samwau ika atapwanapwan, sisap tunano. "
        "Kose mochen akomw pworous ngeni tokter me an noum ewe nenien sukun mwan omw kopwe isoni noum ewe chon sukun non imw."
    )
    for run in footer.runs:
        run.font.size = Pt(10.5)

    # Page break to separate letters
    doc.add_page_break()



def generate_spanish_letter(student, doc, output_dir, school_name, csv_creation_date):
    """
    Generates a Spanish attendance warning letter for a student.
    
    Args:
    - student (dict): Student information.
    - doc (Document): Word document object.
    - output_dir (str): Directory where graph images are stored.
    - school_name (str): Name of the student's school.
    - csv_creation_date (str): Date of the CSV file creation.
    
    Returns:
    - None (adds content to the provided doc object)
    """
    first_name = student['name'].split(', ')[1]
    last_name = student['name'].split(', ')[0]
    student_id = int(float(student['student_number']))  # Ensure integer format

    # Header with attendance warning in Spanish
    header = doc.add_paragraph(f"{first_name} {last_name}, tiene una asistencia inferior al 90% y usted puede ayudar.", style='Heading 1')
    for run in header.runs:
        run.font.size = Pt(12)

    # Student Details
    paragraph = doc.add_paragraph(
        f"Nombre del estudiante: {student['name']}\n"
        f"ID del estudiante: {student_id}\n"
        f"Escuela: {student['school_of_residence']}\n"
        f"Grado: {student['grade']}\n"
        f"Aula de inicio: {student['home_room']}"
    )
    for run in paragraph.runs:
        run.font.size = Pt(10.5)

    # Greeting and introduction in Spanish
    paragraph = doc.add_paragraph(f"Estimado padre o tutor legal de {first_name},")
    for run in paragraph.runs:
        run.font.size = Pt(10.5)
    paragraph = doc.add_paragraph(
        "En el Distrito Escolar CITY, nos esforzamos por mejorar la asistencia de nuestros estudiantes. "
        "Cada minuto de la jornada escolar brinda oportunidades para el crecimiento personal, académico, social y emocional. "
        "Se nota cuando los estudiantes faltan a clase."
    )
    for run in paragraph.runs:
        run.font.size = Pt(10.5)

    # Attendance percentage statement in Spanish
    paragraph = doc.add_paragraph(
        f"{first_name} tiene un porcentaje de {student['current_week_att_percent']}%, sin embargo, el 90% es la expectativa del Distrito y del Estado de Missouri."
    )
    for run in paragraph.runs:
        run.font.size = Pt(10.5)

    # Attendance Graph
    graph_path = create_attendance_graph(f"{first_name} {last_name}", student['current_week_att_percent'], output_dir, school_name, student_id)
    
    doc.add_picture(graph_path, width=Inches(6))
    
    try:
        os.remove(graph_path)
        print(f"Deleted temporary graph: {graph_path}")
    except OSError as e:
        print(f"Error deleting file {graph_path}: {e}")

    # Closing paragraph in Spanish
    paragraph = doc.add_paragraph(
        f"Usted puede tener un gran efecto en mejorar las ausencias de {first_name} este semestre, y agradeceremos su ayuda. "
        f"¡Asegúrese de que {first_name} asista a la escuela a tiempo, todos los días!"
    )
    for run in paragraph.runs:
        run.font.size = Pt(10.5)

    # Signature
    paragraph = doc.add_paragraph(
        f"\nAtentamente,\nEl Equipo de {school_name}\nDistrito Escolar de CITY"
    )
    for run in paragraph.runs:
        run.font.size = Pt(10.5)

    # Footer with initiative details
    paragraph = doc.add_paragraph(
        f"* Este es el porcentaje de asistencia registrado en el distrito escolar de su estudiante hasta la fecha {csv_creation_date}. "
        "Este número incluye ausencias justificadas e injustificadas. Comuníquese directamente con la escuela de su hijo si tiene preguntas sobre las ausencias de su hijo. "
        "Esta carta es parte de la Iniciativa de Asistencia 2025 del School District, que tiene como objetivo aumentar la conciencia sobre la importancia de la asistencia en el School District. "
        "Reconocemos que algunas ausencias, debido a enfermedades o emergencias, son inevitables. Hable con anticipación con la escuela de su estudiante y el proveedor de salud sobre cuándo mantenerlo en casa."
    )
    for run in paragraph.runs:
        run.font.size = Pt(10.5)

    # Page break to separate letters
    doc.add_page_break()


def generate_english_letter(student, doc, output_dir, school_name, csv_creation_date):
    """
    Generates an English attendance warning letter for a student.
    
    Args:
    - student (dict): Student information.
    - doc (Document): Word document object.
    - output_dir (str): Directory where graph images are stored.
    - school_name (str): Name of the student's school.
    - csv_creation_date (str): Date of the CSV file creation.
    
    Returns:
    - None (adds content to the provided doc object)
    """
    first_name = student['name'].split(', ')[1]
    last_name = student['name'].split(', ')[0]
    student_id = int(float(student['student_number']))  # Ensure integer format

    # Header with attendance warning
    header = doc.add_paragraph(f"{first_name} {last_name}, is below 90% attendance and you can help.", style='Heading 1')
    for run in header.runs:
        run.font.size = Pt(12)
    # Student details with line breaks
    paragraph = doc.add_paragraph(
        f"{student['name']}\n"
        f"Student ID: {student_id}\n"
        f"School: {student['school_of_residence']}\n"
        f"Grade: {student['grade']}\n"
        f"Home Room: {student['home_room']}\n"
    )
    for run in paragraph.runs:
        run.font.size = Pt(10.5)

    # Greeting and introduction
    paragraph = doc.add_paragraph(f"Dear Parent/Guardian of {first_name},")
    for run in paragraph.runs:
        run.font.size = Pt(10.5)
    paragraph = doc.add_paragraph(
        "In the School District, we strive to improve the attendance of our students. "
        "Each minute of the school day provides opportunities for personal, academic, social, and emotional growth. "
        "It’s noticeable when students miss time."
    )
    for run in paragraph.runs:
        run.font.size = Pt(10.5)

    # Attendance percentage and expectation statement
    paragraph = doc.add_paragraph(
        f"{first_name} has {student['current_week_att_percent']}% attendance, yet 90% is the District and State of Missouri expectation."
    )
    for run in paragraph.runs:
        run.font.size = Pt(10.5)

    # Create and add the attendance graph
    graph_path = create_attendance_graph(f"{first_name} {last_name}", student['current_week_att_percent'], output_dir, school_name, student_id)
    paragraph = doc.add_paragraph("See the graph below for a visual comparison:")
    for run in paragraph.runs:
        run.font.size = Pt(10.5)
    doc.add_picture(graph_path, width=Inches(6))  # Adjust the width as needed

    # Delete the graph after inserting it into the document
    try:
        os.remove(graph_path)
        #print(f"Deleted temporary graph: {graph_path}")
    except OSError as e:
        print(f"Error deleting file {graph_path}: {e}")

    # Closing paragraph
    paragraph = doc.add_paragraph(
        f"You can have a big effect on {first_name}’s absences this semester - and we appreciate your help. "
        f"Please make sure {first_name} attends school on time, every day!"
    )
    for run in paragraph.runs:
        run.font.size = Pt(10.5)

    # Signature
    paragraph = doc.add_paragraph(
        "\nSincerely,\n"
        f"The {school_name} Team\n"
        "School District"
    )
    for run in paragraph.runs:
        run.font.size = Pt(10.5)

    # Footer with additional initiative details
    paragraph = doc.add_paragraph(
        f"*This is the attendance percentage on record at your student’s school district as of {csv_creation_date}. "
        "This number includes excused and unexcused absences. Please contact your child’s school directly if you have questions "
        "regarding your child’s absences. This letter is part of the School District’s attendance program, which aims to increase awareness "
        "of the importance of attendance in the School District. We recognize that some absences, due to illness or emergencies, are unavoidable. "
        "Please discuss in advance with your student’s school and health provider when to keep your student home."
    )
    for run in paragraph.runs:
        run.font.size = Pt(10.5)

    # Page Break to separate letters
    doc.add_page_break()



def read_final_report(final_report_path):
    """
    Reads the final_report.csv and converts it into a results dictionary.
    
    Args:
    - final_report_path (str): The path to the final_report.csv file.
    
    Returns:
    - dict: A results dictionary with the necessary student information.
    """
    try:
        df = pd.read_csv(final_report_path)
    except FileNotFoundError:
        print(f"Error: {final_report_path} not found.")
        return {}
    
    # Create the results dictionary in the expected format
    results = {}
    for _, row in df.iterrows():
        student_id = str(row['student_number'])
        results[student_id] = {
            'name': row['name'],
            'student_number': student_id,
            'grade': row['grade'],
            'school_of_residence': row['school_of_residence'],
            'attending_school': row['attending_school'],
            'current_week_att_percent': row['current_week_att_percent'],
            'home_room': row['home_room']
        }
    
    return results

def get_csv_creation_date(csv_path):
    """
    Get the creation/modification date of the CSV file and format it as MM/DD/YYYY.
    
    Args:
    - csv_path (str): Path to the CSV file.
    
    Returns:
    - str: The formatted creation date.
    """
    # Get the modification time of the file (last modified date)
    timestamp = os.path.getmtime(csv_path)
    
    # Convert to datetime and format as MM/DD/YYYY
    creation_date = datetime.fromtimestamp(timestamp).strftime('%m/%d/%Y')
    
    return creation_date

def generate_letters_by_attending_school(results, student_languages, base_output_path, csv_creation_date, max_pages=None):
    """
    Generates attendance letters only for students attending alt or specprog.
    Letters are placed in the 'alt and specprog' folder, with testing mode support.

    Args:
    - results (dict): Dictionary of student attendance data.
    - student_languages (dict): Dictionary mapping student IDs to their preferred language.
    - base_output_path (str): Base directory where output folders are created.
    - csv_creation_date (str): Date of the CSV file.
    - max_pages (int, optional): Maximum number of pages per document.

    Returns:
    - None
    """
    # Define the specific schools that qualify for attending school letters
    qualifying_schools = {"alt HS", "alt MS", "alt Elem", "specprog"}

    attending_school_students = {}

    # Filter and group only students attending alt or specprog
    for student_id, student in results.items():
        if student['attending_school'] in qualifying_schools:  # Only process these schools
            attending_school = student['attending_school']
            if attending_school not in attending_school_students:
                attending_school_students[attending_school] = {}
            attending_school_students[attending_school][student_id] = student

    # Adjust output path for testing mode (THIS WAS MISSING)
    if TESTING_MODE and 'test_output' not in base_output_path:
        base_output_path = os.path.join(base_output_path, "test_output")
        print("TESTING MODE: Saving attending school letters to test_output directory.")

    # Create the main alt and specprog directory inside the (potentially modified) output path
    alt_specprog_dir = os.path.join(base_output_path, "alt and specprog")
    os.makedirs(alt_specprog_dir, exist_ok=True)

    # Generate letters only for alt and specprog students
    for attending_school, students in attending_school_students.items():
        school_dir = os.path.join(alt_specprog_dir, attending_school)
        os.makedirs(school_dir, exist_ok=True)

        generate_student_letters(students, student_languages, school_dir, attending_school, csv_creation_date, max_pages)



def main():
    # Path setup (assume the script is run from the directory containing 'output' folder)
    current_dir = os.getcwd()
    output_dir = os.path.join(current_dir, 'output')

    final_report_path = os.path.join(output_dir, 'final_report.csv')
    language_file_path = os.path.join('PS_data.csv')

    # If testing mode is on, set a test output folder
    if TESTING_MODE:
        output_dir = os.path.join(output_dir, 'test_output')
    
    # Get the creation date of the final_report.csv file
    csv_creation_date = get_csv_creation_date(final_report_path)

    building_group_mapping = {
        'Team Lead 1': ['HS 1'],
        'Team Lead 2': ['HS 2'],
        'Team Lead 3': ['MS 3'],
        'Team Lead 4': ['HS 3'],
        'Team Lead 5': ['MS 4', 'ES 1'],
        'Team Lead 6': ['MS 1'],
        'Team Lead 7': ['ES 16'],
        'Team Lead 8': ['ES 13', 'ES 18'],
        'Team Lead 9': ['ES 14', 'ES 15'],
        'Team Lead 10': ['ES 20', 'ES 9'],
        'Team Lead 11': ['ES 21', 'MS 2', 'ES 8', 'ES 10', 'ES 11', 'ES 19'],

        'ES 21': ['ES 21'],
        'MS 2': ['MS 2'],
        'ES 8': ['ES 8'],
        'ES 10': ['ES 10'],
        'ES 11': ['ES 11'],
        'ES 19': ['ES 19'],

        'ES 13': ['ES 13'],

        'Alt ELC': ['ELC'],

        'MS 3': ['MS 3'],
        'ES 1': ['ES 1'],
        'MS 4': ['MS 4'],
        'HS 1': ['HS 1'],
        'ES 15': ['ES 15']
    }


    # Read the final report and recreate the results dictionary
    results = read_final_report(final_report_path)
    student_languages = read_student_languages(language_file_path)

    if not results:
        print("No data found or final_report.csv is missing.")
        return

    # Generate letters by building (now passing student_languages)
    generate_letters_by_building(results, student_languages, output_dir, building_group_mapping, csv_creation_date, 100)

    # Generate letters for alt and specprog by attending school
    generate_letters_by_attending_school(results, student_languages, output_dir, csv_creation_date, 100)

if __name__ == "__main__":
    main()